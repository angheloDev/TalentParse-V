import io
import re
from typing import Any, Dict, List, Optional

from fastapi import FastAPI, File, HTTPException, UploadFile

app = FastAPI(title="Resume Parser Service")

# ---------------------------------------------------------------------------
# Canonical skill list — organised by category. Each value is a list of
# lower-case variants / aliases that count as a match.
# ---------------------------------------------------------------------------
SKILL_SYNONYMS: Dict[str, List[str]] = {
    # Languages
    "JavaScript": ["javascript", "js", "ecmascript", "es6", "es2015", "es2016", "es2017", "es2018", "es2019", "es2020"],
    "TypeScript": ["typescript", "ts"],
    "Python": ["python", "py", "python3", "python2", "python 3"],
    "Java": ["java"],
    "C#": ["c#", "csharp", "c sharp"],
    "C++": ["c++", "cpp", "c plus plus"],
    "C": ["c programming", "ansi c"],
    "Go": ["golang", "go language"],
    "Rust": ["rust", "rustlang"],
    "Swift": ["swift", "swiftlang"],
    "Kotlin": ["kotlin"],
    "PHP": ["php"],
    "Ruby": ["ruby"],
    "Scala": ["scala"],
    "Dart": ["dart"],
    "Elixir": ["elixir"],
    "Haskell": ["haskell"],
    "Lua": ["lua"],
    "Objective-C": ["objective-c", "objc", "objective c"],
    "Bash": ["bash", "shell script", "shell scripting"],
    "PowerShell": ["powershell", "pwsh"],
    "Groovy": ["groovy"],
    "Perl": ["perl"],
    "R": ["r programming", "r language"],
    "MATLAB": ["matlab"],
    # Frontend
    "React": ["react", "reactjs", "react.js"],
    "Vue.js": ["vue", "vuejs", "vue.js", "vue 3", "vue3"],
    "Angular": ["angular", "angularjs"],
    "Svelte": ["svelte", "sveltekit"],
    "Next.js": ["next.js", "nextjs", "next js"],
    "Nuxt.js": ["nuxt", "nuxt.js", "nuxtjs"],
    "Gatsby": ["gatsby"],
    "HTML": ["html", "html5"],
    "CSS": ["css", "css3"],
    "Sass": ["sass", "scss"],
    "Tailwind CSS": ["tailwind", "tailwindcss", "tailwind css"],
    "Bootstrap": ["bootstrap"],
    "Material UI": ["material ui", "mui", "material-ui"],
    "Redux": ["redux", "redux toolkit"],
    "Zustand": ["zustand"],
    "Webpack": ["webpack"],
    "Vite": ["vite", "vitejs"],
    "jQuery": ["jquery"],
    "Storybook": ["storybook"],
    "Cypress": ["cypress"],
    "Playwright": ["playwright"],
    "Selenium": ["selenium"],
    "Jest": ["jest"],
    # Backend
    "Node.js": ["node.js", "nodejs", "node js"],
    "Express": ["express", "express.js", "expressjs"],
    "NestJS": ["nestjs", "nest.js"],
    "Django": ["django"],
    "Flask": ["flask"],
    "FastAPI": ["fastapi"],
    "Spring Boot": ["spring boot", "springboot"],
    "Spring": ["spring framework", "spring mvc"],
    "Laravel": ["laravel"],
    "Ruby on Rails": ["rails", "ruby on rails", "ror"],
    "ASP.NET": ["asp.net", "aspnet", "asp.net core"],
    "Gin": ["gin-gonic", "gin gonic"],
    "Fastify": ["fastify"],
    "Koa": ["koa.js"],
    # Databases
    "PostgreSQL": ["postgresql", "postgres", "pgsql"],
    "MySQL": ["mysql"],
    "MongoDB": ["mongodb", "mongo"],
    "Redis": ["redis"],
    "Elasticsearch": ["elasticsearch", "elastic search"],
    "Cassandra": ["cassandra", "apache cassandra"],
    "SQLite": ["sqlite"],
    "Oracle Database": ["oracle database", "oracle db"],
    "SQL Server": ["sql server", "mssql"],
    "DynamoDB": ["dynamodb", "dynamo db"],
    "Firebase": ["firebase", "firestore"],
    "Supabase": ["supabase"],
    "Neo4j": ["neo4j"],
    "Snowflake": ["snowflake"],
    "BigQuery": ["bigquery"],
    "Amazon Redshift": ["redshift"],
    "Databricks": ["databricks"],
    "RabbitMQ": ["rabbitmq"],
    "Apache Kafka": ["kafka", "apache kafka"],
    "Prisma": ["prisma"],
    "TypeORM": ["typeorm"],
    "Sequelize": ["sequelize"],
    "Mongoose": ["mongoose"],
    # Cloud & DevOps
    "AWS": ["aws", "amazon web services"],
    "Microsoft Azure": ["azure", "microsoft azure"],
    "Google Cloud": ["gcp", "google cloud", "google cloud platform"],
    "Docker": ["docker"],
    "Kubernetes": ["kubernetes", "k8s"],
    "Terraform": ["terraform"],
    "Ansible": ["ansible"],
    "Helm": ["helm"],
    "Prometheus": ["prometheus"],
    "Grafana": ["grafana"],
    "Datadog": ["datadog"],
    "New Relic": ["new relic", "newrelic"],
    "Cloudflare": ["cloudflare"],
    "Vercel": ["vercel"],
    "Netlify": ["netlify"],
    "Heroku": ["heroku"],
    "Linux": ["linux", "ubuntu", "centos", "debian"],
    "Nginx": ["nginx"],
    "Serverless": ["serverless"],
    "Microservices": ["microservices", "microservice"],
    # CI/CD
    "Git": ["git"],
    "GitHub": ["github"],
    "GitLab": ["gitlab"],
    "GitHub Actions": ["github actions"],
    "Jenkins": ["jenkins"],
    "CircleCI": ["circleci"],
    # Data & ML
    "TensorFlow": ["tensorflow", "tf"],
    "PyTorch": ["pytorch", "torch"],
    "scikit-learn": ["scikit-learn", "sklearn", "scikit learn"],
    "Pandas": ["pandas"],
    "NumPy": ["numpy"],
    "Apache Spark": ["pyspark", "apache spark"],
    "Apache Airflow": ["airflow", "apache airflow"],
    "dbt": ["dbt"],
    "Power BI": ["power bi", "powerbi"],
    "Tableau": ["tableau"],
    "Machine Learning": ["machine learning"],
    "Deep Learning": ["deep learning", "neural network", "neural networks"],
    "NLP": ["nlp", "natural language processing"],
    "Computer Vision": ["computer vision", "image recognition"],
    "LangChain": ["langchain"],
    # Mobile
    "React Native": ["react native", "react-native"],
    "Flutter": ["flutter"],
    "SwiftUI": ["swiftui", "swift ui"],
    "Jetpack Compose": ["jetpack compose"],
    "Ionic": ["ionic"],
    "Expo": ["expo"],
    # APIs & protocols
    "GraphQL": ["graphql", "gql"],
    "REST API": ["rest api", "restful", "rest apis"],
    "gRPC": ["grpc"],
    "WebSocket": ["websocket", "websockets"],
    "OAuth": ["oauth", "oauth2"],
    "JWT": ["jwt", "json web token"],
    "OpenAPI": ["openapi", "swagger"],
    # Testing
    "Jest": ["jest"],
    "pytest": ["pytest"],
    "JUnit": ["junit"],
    "Cypress": ["cypress"],
    "TDD": ["tdd", "test driven development"],
    # Methodologies
    "Agile": ["agile"],
    "Scrum": ["scrum"],
    "Kanban": ["kanban"],
    "DevOps": ["devops"],
    "CI/CD": ["ci/cd", "cicd", "continuous integration", "continuous deployment"],
    # Tools
    "Jira": ["jira"],
    "Figma": ["figma"],
    "Postman": ["postman"],
    # Security
    "Cybersecurity": ["cybersecurity", "cyber security"],
    "Penetration Testing": ["penetration testing", "pentesting"],
    "OWASP": ["owasp"],
}

CERTIFICATION_PATTERNS: List[re.Pattern] = [
    re.compile(r'AWS\s+Certified[^.\n]{0,80}', re.IGNORECASE),
    re.compile(r'Google\s+(?:Certified|Professional|Associate)[^.\n]{0,60}', re.IGNORECASE),
    re.compile(r'Microsoft\s+Certified[^.\n]{0,60}', re.IGNORECASE),
    re.compile(r'Oracle\s+Certified[^.\n]{0,60}', re.IGNORECASE),
    re.compile(r'Certified\s+(?:Kubernetes|ScrumMaster|Scrum\s+Master|Product\s+Owner|Data\s+Professional|Cloud\s+Engineer)[^.\n]{0,60}', re.IGNORECASE),
    re.compile(r'\b(?:PMP|CISSP|CEH|CISM|CISA|CPA|CFA|TOGAF)\b[^.\n]{0,40}', re.IGNORECASE),
    re.compile(r'CompTIA\s+(?:A\+|Network\+|Security\+|Cloud\+|CySA\+|CASP\+)[^.\n]{0,40}', re.IGNORECASE),
    re.compile(r'Certified\s+(?:Ethical\s+Hacker|Information\s+Systems|Security\s+Professional)[^.\n]{0,60}', re.IGNORECASE),
]

DEGREE_PATTERNS: List[re.Pattern] = [
    re.compile(r'\b(Ph\.?D\.?|Doctor(?:ate)?(?:\s+of\s+\w+)?)[^.\n]{0,80}', re.IGNORECASE),
    re.compile(r'\b(Master(?:s)?(?:\s+of\s+\w+)?|M\.?S\.?|M\.?A\.?|MBA|M\.?Eng\.?|M\.?Tech\.?|M\.?Sc\.?)[^.\n]{0,80}', re.IGNORECASE),
    re.compile(r'\b(Bachelor(?:s)?(?:\s+of\s+\w+)?|B\.?S\.?|B\.?A\.?|B\.?Eng\.?|B\.?Tech\.?|B\.?Sc\.?)[^.\n]{0,80}', re.IGNORECASE),
    re.compile(r'\b(Associate(?:s)?(?:\s+of\s+\w+)?|A\.?S\.?|A\.?A\.?)[^.\n]{0,60}', re.IGNORECASE),
    re.compile(r'\b(High\s+School\s+Diploma|GED|Diploma)[^.\n]{0,60}', re.IGNORECASE),
]


def _escape_re(s: str) -> str:
    return re.escape(s)


def _extract_text_pdf(content: bytes) -> str:
    import pdfplumber
    parts: List[str] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            text = page.extract_text(x_tolerance=2, y_tolerance=2)
            if text:
                parts.append(text)
    return "\n".join(parts)


def _extract_text_docx(content: bytes) -> str:
    import docx
    doc = docx.Document(io.BytesIO(content))
    paragraphs: List[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                stripped = cell.text.strip()
                if stripped:
                    paragraphs.append(stripped)
    return "\n".join(paragraphs)


def extract_email(text: str) -> str:
    m = re.search(r'\b[A-Z0-9._%+\-]+@[A-Z0-9.\-]+\.[A-Z]{2,}\b', text, re.IGNORECASE)
    return m.group(0) if m else ""


def extract_phone(text: str) -> str:
    m = re.search(r'(?:\+?\d[\d\s\-().]{7,20}\d)', text)
    return re.sub(r'\s+', ' ', m.group(0)).strip() if m else ""


def extract_linkedin(text: str) -> str:
    m = re.search(r'linkedin\.com/in/([A-Za-z0-9\-_%]+)', text, re.IGNORECASE)
    if m:
        return f"https://linkedin.com/in/{m.group(1)}"
    return ""


def extract_github(text: str) -> str:
    m = re.search(r'github\.com/([A-Za-z0-9\-]+)(?:/[A-Za-z0-9\-\.]+)?', text, re.IGNORECASE)
    if m:
        return f"https://github.com/{m.group(1)}"
    return ""


def extract_portfolio(text: str) -> str:
    # URLs that are not LinkedIn or GitHub
    matches = re.findall(
        r'https?://(?!(?:www\.)?(?:linkedin|github)\.com)[^\s<>"\'()]+',
        text,
        re.IGNORECASE,
    )
    return matches[0].rstrip('.,;') if matches else ""


def extract_name(text: str) -> str:
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    for line in lines[:20]:
        if len(line) < 4 or len(line) > 65:
            continue
        if re.search(r'[@|•\d]|http|linkedin|github|resume|curriculum|vitae|email|phone|address', line, re.IGNORECASE):
            continue
        # Matches 2-4 capitalised name parts
        if re.match(r'^[A-Z][a-zA-Z\'\-]+(?:\s+[A-Z][a-zA-Z\'\-]+){1,3}$', line):
            return line
    return ""


def extract_skills(text: str) -> List[str]:
    found: Dict[str, bool] = {}
    text_lower = f" {text.lower()} "

    for canonical, variants in SKILL_SYNONYMS.items():
        for variant in variants:
            pattern = r'(?<![a-z0-9])' + _escape_re(variant) + r'(?![a-z0-9])'
            if re.search(pattern, text_lower):
                found[canonical] = True
                break

    # Also scrape skills listed in skills sections for items not in our dict
    section_match = re.search(
        r'(?:technical\s+)?skills?(?:\s+&\s+\w+)?(?:\s+/\s+\w+)?\s*[:\-]\s*([^\n]{5,400})',
        text,
        re.IGNORECASE,
    )
    if section_match:
        raw_items = re.split(r'[,|•/·\t]+', section_match.group(1))
        for item in raw_items:
            item = item.strip().strip('•·-').strip()
            if 2 <= len(item) <= 50:
                item_lower = item.lower()
                for canonical, variants in SKILL_SYNONYMS.items():
                    if canonical not in found:
                        if any(v in item_lower or item_lower in v for v in variants):
                            found[canonical] = True
                            break

    return sorted(found.keys())


def extract_certifications(text: str) -> List[Dict[str, str]]:
    certs: List[Dict[str, str]] = []
    seen: set = set()
    for pattern in CERTIFICATION_PATTERNS:
        for m in pattern.finditer(text):
            raw = m.group(0).strip()
            key = raw.lower()
            if key not in seen and 4 < len(raw) < 200:
                seen.add(key)
                year_m = re.search(r'\b(20\d{2})\b', raw)
                name = re.sub(r'\b20\d{2}\b', '', raw).strip(' ,;')
                certs.append({"name": name, "issuer": "", "year": year_m.group(1) if year_m else ""})
    return certs


def extract_education(text: str) -> str:
    degrees: List[str] = []
    seen: set = set()
    for pattern in DEGREE_PATTERNS:
        for m in pattern.finditer(text):
            raw = m.group(0).strip()
            key = raw.lower()[:30]
            if key not in seen:
                seen.add(key)
                degrees.append(raw[:120])
        if degrees:
            break  # stop at highest degree found
    return ", ".join(degrees) if degrees else ""


def extract_total_experience(text: str) -> Optional[str]:
    m = re.search(
        r'(\d+(?:\.\d+)?)\+?\s*years?\s+(?:of\s+)?(?:professional\s+)?experience',
        text,
        re.IGNORECASE,
    )
    if m:
        return f"{m.group(1)} years"
    return None


@app.post("/parse-resume")
async def parse_resume(file: UploadFile = File(...)) -> Dict[str, Any]:
    content_type = (file.content_type or "").lower()
    fname = (file.filename or "").lower()

    is_pdf = "pdf" in content_type or fname.endswith(".pdf")
    is_docx = "wordprocessingml" in content_type or fname.endswith(".docx")

    if not is_pdf and not is_docx:
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported.")

    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    try:
        text = _extract_text_pdf(content) if is_pdf else _extract_text_docx(content)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Text extraction failed: {exc}") from exc

    if not text.strip():
        raise HTTPException(status_code=422, detail="Could not extract any text from the file.")

    skills = extract_skills(text)
    experience = extract_total_experience(text) or ""

    return {
        "name": extract_name(text),
        "email": extract_email(text),
        "phone": extract_phone(text),
        "linkedin": extract_linkedin(text),
        "github": extract_github(text),
        "portfolio": extract_portfolio(text),
        "skills": skills,
        "experience": experience,
        "education": extract_education(text),
        "certifications": extract_certifications(text),
        "rawText": text,
    }
